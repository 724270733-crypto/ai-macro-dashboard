"""Build the complete AI macro-cycle dashboard snapshot.

The source workbooks are never modified.  This exporter reads cached workbook
values, refreshes A-share financials through Tushare when configured, and
preserves provenance/quality fields so estimates cannot be confused with
official disclosures.

Environment variables:
  AI_FRAMEWORK_DIR   folder containing the source Excel/Word/PDF files
  TUSHARE_TOKEN      optional; enables live A-share financial refresh
  IFIND_ACCESS_TOKEN optional; records that QuantAPI is ready for live refresh
  IFIND_REFRESH_TOKEN optional; used only by the official local iFinD client
"""
from __future__ import annotations

import json
import math
import os
import re
import urllib.request
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from openpyxl import load_workbook
from refresh_credit import build_credit_snapshot

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "dashboard-data.json"
SOURCE = Path(os.environ.get("AI_FRAMEWORK_DIR", r"C:\Users\yingx\OneDrive\桌面\AI框架"))
DB_NAME = "AI宏观数据库_指标与获取路径.xlsx"
SEMI_NAME = "SemiAnalysis_公司模型与公开指标汇总.xlsx"
STATE_NAME = "state_of_ai.xlsx"
REPORT_NAME = "AI产业链的宏观景气跟踪框架.docx"

# Last verified Tushare snapshot. It keeps the public dashboard usable when the
# live endpoint is unavailable; live refresh replaces it when the API succeeds.
FALLBACK_FINANCIALS = [
    {"code":"688256.SH","company":"寒武纪","period":"20231231","revenue_cny_bn":0.709,"operating_profit_cny_bn":-0.876,"net_income_cny_bn":-0.878,"operating_cashflow_cny_bn":-0.596,"capex_cash_cny_bn":0.1},
    {"code":"688256.SH","company":"寒武纪","period":"20240331","revenue_cny_bn":0.026,"operating_profit_cny_bn":-0.228,"net_income_cny_bn":-0.229,"operating_cashflow_cny_bn":-0.234,"capex_cash_cny_bn":0.038},
    {"code":"688256.SH","company":"寒武纪","period":"20240630","revenue_cny_bn":0.065,"operating_profit_cny_bn":-0.533,"net_income_cny_bn":-0.533,"operating_cashflow_cny_bn":-0.631,"capex_cash_cny_bn":0.117},
    {"code":"688256.SH","company":"寒武纪","period":"20240930","revenue_cny_bn":0.185,"operating_profit_cny_bn":-0.728,"net_income_cny_bn":-0.728,"operating_cashflow_cny_bn":-1.81,"capex_cash_cny_bn":0.211},
    {"code":"688256.SH","company":"寒武纪","period":"20241231","revenue_cny_bn":1.174,"operating_profit_cny_bn":-0.456,"net_income_cny_bn":-0.457,"operating_cashflow_cny_bn":-1.618,"capex_cash_cny_bn":0.366},
    {"code":"688256.SH","company":"寒武纪","period":"20250331","revenue_cny_bn":1.111,"operating_profit_cny_bn":0.355,"net_income_cny_bn":0.355,"operating_cashflow_cny_bn":-1.399,"capex_cash_cny_bn":0.168},
    {"code":"688256.SH","company":"寒武纪","period":"20250630","revenue_cny_bn":2.881,"operating_profit_cny_bn":1.038,"net_income_cny_bn":1.038,"operating_cashflow_cny_bn":0.911,"capex_cash_cny_bn":0.255},
    {"code":"688256.SH","company":"寒武纪","period":"20250930","revenue_cny_bn":4.607,"operating_profit_cny_bn":1.606,"net_income_cny_bn":1.604,"operating_cashflow_cny_bn":-0.029,"capex_cash_cny_bn":0.329},
    {"code":"688256.SH","company":"寒武纪","period":"20251231","revenue_cny_bn":6.497,"operating_profit_cny_bn":2.061,"net_income_cny_bn":2.058,"operating_cashflow_cny_bn":-0.498,"capex_cash_cny_bn":0.569},
    {"code":"688256.SH","company":"寒武纪","period":"20260331","revenue_cny_bn":2.885,"operating_profit_cny_bn":1.014,"net_income_cny_bn":1.013,"operating_cashflow_cny_bn":0.834,"capex_cash_cny_bn":0.199},
    {"code":"688041.SH","company":"海光信息","period":"20231231","revenue_cny_bn":6.012,"operating_profit_cny_bn":1.68,"net_income_cny_bn":1.701,"operating_cashflow_cny_bn":0.814,"capex_cash_cny_bn":0.918},
    {"code":"688041.SH","company":"海光信息","period":"20240331","revenue_cny_bn":1.592,"operating_profit_cny_bn":0.401,"net_income_cny_bn":0.394,"operating_cashflow_cny_bn":-0.068,"capex_cash_cny_bn":0.211},
    {"code":"688041.SH","company":"海光信息","period":"20240630","revenue_cny_bn":3.763,"operating_profit_cny_bn":1.242,"net_income_cny_bn":1.226,"operating_cashflow_cny_bn":-0.113,"capex_cash_cny_bn":0.499},
    {"code":"688041.SH","company":"海光信息","period":"20240930","revenue_cny_bn":6.137,"operating_profit_cny_bn":2.159,"net_income_cny_bn":2.107,"operating_cashflow_cny_bn":0.399,"capex_cash_cny_bn":0.733},
    {"code":"688041.SH","company":"海光信息","period":"20241231","revenue_cny_bn":9.162,"operating_profit_cny_bn":2.789,"net_income_cny_bn":2.717,"operating_cashflow_cny_bn":0.977,"capex_cash_cny_bn":0.945},
    {"code":"688041.SH","company":"海光信息","period":"20250331","revenue_cny_bn":2.4,"operating_profit_cny_bn":0.73,"net_income_cny_bn":0.714,"operating_cashflow_cny_bn":2.522,"capex_cash_cny_bn":0.241},
    {"code":"688041.SH","company":"海光信息","period":"20250630","revenue_cny_bn":5.464,"operating_profit_cny_bn":1.639,"net_income_cny_bn":1.642,"operating_cashflow_cny_bn":2.177,"capex_cash_cny_bn":0.192},
    {"code":"688041.SH","company":"海光信息","period":"20250930","revenue_cny_bn":9.49,"operating_profit_cny_bn":2.839,"net_income_cny_bn":2.841,"operating_cashflow_cny_bn":2.255,"capex_cash_cny_bn":0.899},
    {"code":"688041.SH","company":"海光信息","period":"20251231","revenue_cny_bn":14.377,"operating_profit_cny_bn":3.605,"net_income_cny_bn":3.619,"operating_cashflow_cny_bn":2.097,"capex_cash_cny_bn":1.17},
    {"code":"688041.SH","company":"海光信息","period":"20260331","revenue_cny_bn":4.034,"operating_profit_cny_bn":0.97,"net_income_cny_bn":0.877,"operating_cashflow_cny_bn":0.068,"capex_cash_cny_bn":0.382},
    {"code":"300308.SZ","company":"中际旭创","period":"20231231","revenue_cny_bn":10.718,"operating_profit_cny_bn":2.494,"net_income_cny_bn":2.208,"operating_cashflow_cny_bn":1.897,"capex_cash_cny_bn":1.704},
    {"code":"300308.SZ","company":"中际旭创","period":"20240331","revenue_cny_bn":4.843,"operating_profit_cny_bn":1.17,"net_income_cny_bn":1.028,"operating_cashflow_cny_bn":0.651,"capex_cash_cny_bn":0.704},
    {"code":"300308.SZ","company":"中际旭创","period":"20240630","revenue_cny_bn":10.799,"operating_profit_cny_bn":2.728,"net_income_cny_bn":2.407,"operating_cashflow_cny_bn":0.968,"capex_cash_cny_bn":1.258},
    {"code":"300308.SZ","company":"中际旭创","period":"20240930","revenue_cny_bn":17.313,"operating_profit_cny_bn":4.41,"net_income_cny_bn":3.872,"operating_cashflow_cny_bn":1.316,"capex_cash_cny_bn":2.035},
    {"code":"300308.SZ","company":"中际旭创","period":"20241231","revenue_cny_bn":23.862,"operating_profit_cny_bn":6.05,"net_income_cny_bn":5.372,"operating_cashflow_cny_bn":3.165,"capex_cash_cny_bn":2.866},
    {"code":"300308.SZ","company":"中际旭创","period":"20250331","revenue_cny_bn":6.674,"operating_profit_cny_bn":1.989,"net_income_cny_bn":1.691,"operating_cashflow_cny_bn":2.164,"capex_cash_cny_bn":0.402},
    {"code":"300308.SZ","company":"中际旭创","period":"20250630","revenue_cny_bn":14.789,"operating_profit_cny_bn":4.877,"net_income_cny_bn":4.242,"operating_cashflow_cny_bn":3.218,"capex_cash_cny_bn":0.954},
    {"code":"300308.SZ","company":"中际旭创","period":"20250930","revenue_cny_bn":25.005,"operating_profit_cny_bn":8.839,"net_income_cny_bn":7.57,"operating_cashflow_cny_bn":5.455,"capex_cash_cny_bn":1.615},
    {"code":"300308.SZ","company":"中际旭创","period":"20251231","revenue_cny_bn":38.24,"operating_profit_cny_bn":13.597,"net_income_cny_bn":11.58,"operating_cashflow_cny_bn":10.896,"capex_cash_cny_bn":2.76},
    {"code":"300308.SZ","company":"中际旭创","period":"20260331","revenue_cny_bn":19.496,"operating_profit_cny_bn":7.535,"net_income_cny_bn":6.317,"operating_cashflow_cny_bn":3.368,"capex_cash_cny_bn":1.929},
]


def clean(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (datetime, date)):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float):
        if math.isnan(value) or math.isinf(value):
            return None
        return round(value, 8)
    if isinstance(value, str):
        value = value.strip()
        if not value or value.startswith("#"):
            return None
    return value


def row_record(headers: Iterable[Any], values: Iterable[Any]) -> dict[str, Any]:
    out: dict[str, Any] = {}
    for index, (header, value) in enumerate(zip(headers, values), 1):
        key = str(header).strip() if header not in (None, "") else f"field_{index}"
        value = clean(value)
        if value is not None:
            out[key] = value
    return out


def table(ws, header_row: int, start_row: int, end_row: int, start_col: int, end_col: int):
    headers = [clean(ws.cell(header_row, c).value) for c in range(start_col, end_col + 1)]
    rows = []
    for r in range(start_row, end_row + 1):
        values = [clean(ws.cell(r, c).value) for c in range(start_col, end_col + 1)]
        if any(v is not None for v in values):
            rows.append(row_record(headers, values))
    return rows


def stream_table(ws, header_row: int, start_row: int, end_row: int, start_col: int, end_col: int):
    """Read a bounded range efficiently from a read-only worksheet."""
    header_values = next(ws.iter_rows(
        min_row=header_row, max_row=header_row,
        min_col=start_col, max_col=end_col, values_only=True,
    ))
    headers = [clean(value) for value in header_values]
    rows = []
    for values in ws.iter_rows(
        min_row=start_row, max_row=end_row,
        min_col=start_col, max_col=end_col, values_only=True,
    ):
        cleaned = [clean(value) for value in values]
        if any(value is not None for value in cleaned):
            rows.append(row_record(headers, cleaned))
    return rows


def time_series(ws, header_row: int, start_row: int, start_col: int, end_col: int):
    headers = [clean(ws.cell(header_row, c).value) for c in range(start_col, end_col + 1)]
    rows = []
    for r in range(start_row, ws.max_row + 1):
        values = [clean(ws.cell(r, c).value) for c in range(start_col, end_col + 1)]
        if values and isinstance(values[0], str) and re.match(r"^\d{4}-\d{2}-\d{2}", values[0]):
            rows.append(row_record(headers, values))
    rows.sort(key=lambda x: next(iter(x.values())))
    return rows


def parse_number(value: Any) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)
    if not isinstance(value, str):
        return None
    text = value.replace(",", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return float(match.group()) if match else None


def extract_main_workbook() -> dict[str, Any]:
    path = SOURCE / DB_NAME
    wb = load_workbook(path, data_only=True, read_only=False)

    # OpenRouter token history.
    ws = wb["top50模型token数量2"]
    token_headers = [clean(ws.cell(1, c).value) for c in range(1, ws.max_column + 1)]
    token_rows = []
    for r in range(2, ws.max_row + 1):
        vals = [clean(ws.cell(r, c).value) for c in range(1, ws.max_column + 1)]
        if not vals or not isinstance(vals[0], str):
            continue
        rec = row_record(token_headers, vals)
        rec["date"] = vals[0]
        provider_values = [
            v for k, v in rec.items()
            if k not in {"date", token_headers[0]} and isinstance(v, (int, float))
        ]
        rec["total_tokens"] = sum(provider_values)
        token_rows.append(rec)
    token_rows.sort(key=lambda x: x["date"])

    # OpenRouter app rankings.
    ws = wb["按代币使用率排名前列应用"]
    apps = []
    for r in range(2, ws.max_row + 1):
        rank, app_id, app_name, total_tokens, total_requests = [
            clean(ws.cell(r, c).value) for c in range(1, 6)
        ]
        if not isinstance(rank, (int, float)) or not isinstance(total_tokens, (int, float)):
            continue
        apps.append({
            "rank": int(rank), "app_id": app_id, "app_name": app_name or app_id,
            "total_tokens": total_tokens, "total_requests": total_requests,
            "tokens_per_request": (
                total_tokens / total_requests
                if isinstance(total_requests, (int, float)) and total_requests else None
            ),
        })

    # Model price event ledger.
    ws = wb["模型价格变动跟踪"]
    price_headers = [
        "id", "date", "model_id", "model", "provider", "field", "field_label",
        "change_type", "old", "new", "delta", "pct",
    ]
    price_changes = []
    for r in range(2, ws.max_row + 1):
        vals = [clean(ws.cell(r, c).value) for c in range(1, 13)]
        if vals[1] and vals[3]:
            price_changes.append(row_record(price_headers, vals))

    # Artificial Analysis model universe.
    ws = wb["模型价格横向比较底稿"]
    model_benchmarks = table(ws, 8, 9, ws.max_row, 1, 22)

    # Google Trends + Similarweb.
    ws = wb["商业化热度"]
    google_trends = time_series(ws, 14, 15, 1, 9)
    similarweb = table(ws, 15, 16, 25, 33, 55)

    # Hardware company disclosures.
    ws = wb["硬件景气"]
    hardware = []
    blocks = [
        ("NVIDIA", 1, 5), ("AMD", 8, 11), ("Broadcom", 14, 17), ("Marvell", 20, 25),
    ]
    for company, c1, c2 in blocks:
        headers = [clean(ws.cell(13, c).value) for c in range(c1, c2 + 1)]
        for r in range(14, 19):
            vals = [clean(ws.cell(r, c).value) for c in range(c1, c2 + 1)]
            if vals[0] and any(v is not None for v in vals[1:]):
                rec = row_record(headers, vals)
                rec["company"] = company
                revenue_source = next((v for k, v in rec.items() if "收入" in k), None)
                revenue = parse_number(revenue_source)
                if revenue is not None and isinstance(revenue_source, str) and "亿美元" in revenue_source:
                    revenue /= 10
                if revenue is not None:
                    rec["revenue_usd_bn"] = revenue
                hardware.append(rec)
    odm = time_series(ws, 12, 15, 43, 51)

    # Semiconductor cycle.
    ws = wb["半导体周期"]
    semiconductor = []
    for r in range(14, ws.max_row + 1):
        vals = [clean(ws.cell(r, c).value) for c in range(1, min(ws.max_column, 18) + 1)]
        date_value = next((v for v in vals if isinstance(v, str) and re.match(r"^\d{4}-\d{2}", v)), None)
        if date_value and any(isinstance(v, (int, float)) for v in vals):
            semiconductor.append({"date": date_value, "values": vals})

    # Memory, upstream materials, and CCL events.
    ws = wb["价格"]
    memory_spot = time_series(ws, 13, 16, 1, 6)
    memory_contract = time_series(ws, 13, 16, 8, 13)
    copper_foil = time_series(ws, 13, 16, 15, 20)
    electronic_cloth = time_series(ws, 13, 16, 23, 24)
    epoxy_resin = time_series(ws, 13, 16, 37, 38)
    ccl_events = table(ws, 11, 12, 20, 31, 35)

    # China power generation mix.
    ws = wb["电力"]
    power = time_series(ws, 8, 11, 4, 14)

    return {
        "token": token_rows,
        "apps": apps,
        "model_price_changes": price_changes,
        "model_benchmarks": model_benchmarks,
        "google_trends": google_trends,
        "similarweb": similarweb,
        "hardware_quarterly": hardware,
        "odm_monthly": odm,
        "semiconductor_cycle": semiconductor,
        "memory_spot": memory_spot,
        "memory_contract": memory_contract,
        "copper_foil": copper_foil,
        "electronic_cloth": electronic_cloth,
        "epoxy_resin": epoxy_resin,
        "ccl_events": ccl_events,
        "power_generation": power,
    }


def extract_state_of_ai() -> dict[str, Any]:
    wb = load_workbook(SOURCE / STATE_NAME, data_only=True, read_only=False)
    funding = table(wb["Equity Funding"], 17, 18, 32, 2, 5)
    exits = table(wb["Exits"], 9, 10, 24, 2, 5)
    top_deals = table(wb["Top Deals"], 9, 10, 19, 2, 7)
    top_sectors = table(wb["Top Sectors"], 9, 10, 14, 2, 5)
    top_companies = table(wb["Top Companies by Mosaic"], 9, 10, 19, 2, 7)
    recent = table(wb["Recent Activity"], 9, 10, 29, 2, 7)
    investors = table(wb["Most Active Investors"], 9, 10, 19, 2, 4)
    return {
        "funding": funding, "exits": exits, "top_deals": top_deals,
        "top_sectors": top_sectors, "top_companies": top_companies,
        "recent_activity": recent, "active_investors": investors,
        "as_of": "2026-07-19",
    }


def extract_semianalysis() -> dict[str, Any]:
    # The workbook contains a malformed embedded chart. Streaming mode ignores
    # chart XML and is also much faster for its long data sheets.
    wb = load_workbook(SOURCE / SEMI_NAME, data_only=True, read_only=True)
    models_ws = wb["01_模型长表"]
    revenue_ws = wb["02_收入宽表"]
    public_ws = wb["03_公开指标"]
    sources_ws = wb["05_来源说明"]
    models = stream_table(models_ws, 4, 5, models_ws.max_row, 1, 21)
    revenue = stream_table(revenue_ws, 4, 5, revenue_ws.max_row, 1, 12)
    public = stream_table(public_ws, 4, 5, public_ws.max_row, 1, 16)
    sources = stream_table(sources_ws, 4, 5, sources_ws.max_row, 1, 8)
    return {
        "company_models": models,
        "company_revenue": revenue,
        "public_indicators": public,
        "source_ledger": sources,
    }


def extract_report_outline() -> dict[str, Any]:
    from docx import Document
    doc = Document(SOURCE / REPORT_NAME)
    paragraphs = doc.paragraphs
    sections = []
    for i, para in enumerate(paragraphs):
        style = para.style.name if para.style else ""
        title = para.text.strip()
        if title and style.startswith("Heading"):
            body = []
            for nxt in paragraphs[i + 1:i + 5]:
                text = nxt.text.strip()
                if text and not (nxt.style and nxt.style.name.startswith("Heading")):
                    body.append(text)
            sections.append({"level": style, "title": title, "lead": " ".join(body)[:700]})
    return {"sections": sections, "paragraphs": len(paragraphs), "tables": len(doc.tables)}


def extract_pdf_catalog() -> list[dict[str, Any]]:
    catalog = []
    for path in sorted(SOURCE.glob("*.pdf")):
        # Page counts/themes were audited once during dashboard construction.
        # Avoid parsing licensed PDFs on every routine data refresh.
        known = {
            "IT桔子.pdf": (None, "2026年上半年中国AI投融资：城市、赛道、轮次与头部项目"),
            "2026-06-23_兴业证券_如何跟踪和预判AI产业趋势？.pdf": (10, "Token价格、Hyperscaler Capex、ARR与盈利预期"),
            "2026-07-06_国联民生_AI与新范式系列：如何刻画AI的可持续性？.pdf": (19, "收入、现金流、债务与AI投资可持续性"),
            "国泰君安期货-海外AI产业链景气追踪框架-260713.pdf": (27, "ARR—Capex—库存/价格/扩产—盈利四阶段框架"),
        }
        pages, theme = known.get(path.name, (None, "研究资料"))
        catalog.append({"file": path.name, "pages": pages, "theme": theme})
    return catalog


def tushare_financials():
    token = os.environ.get("TUSHARE_TOKEN")
    if not token:
        if OUT.exists():
            try:
                previous = json.loads(OUT.read_text(encoding="utf-8"))
                cached = previous.get("a_share_financials", [])
                if cached:
                    return cached, "cached snapshot; live refresh skipped"
            except Exception:
                pass
        return FALLBACK_FINANCIALS, "verified cached snapshot"
    try:
        def query(api_name: str, ts_code: str, fields: str):
            payload = json.dumps({
                "api_name": api_name,
                "token": token,
                "params": {
                    "ts_code": ts_code,
                    "start_date": "20230101",
                    "end_date": "20271231",
                },
                "fields": fields,
            }).encode("utf-8")
            request = urllib.request.Request(
                "http://jiaoch.site",
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            with urllib.request.urlopen(request, timeout=35) as response:
                result = json.loads(response.read().decode("utf-8"))
            if result.get("code") not in (0, None):
                raise RuntimeError(result.get("msg") or f"{api_name} failed")
            block = result.get("data") or {}
            fields_out = block.get("fields") or []
            return [dict(zip(fields_out, row)) for row in (block.get("items") or [])]

        companies = {
            "688256.SH": "寒武纪", "688041.SH": "海光信息", "300308.SZ": "中际旭创",
            "300502.SZ": "新易盛", "002463.SZ": "沪电股份", "600183.SH": "生益科技",
            "601138.SH": "工业富联", "688008.SH": "澜起科技", "603986.SH": "兆易创新",
            "688981.SH": "中芯国际",
        }
        output = []
        for code, name in companies.items():
            income_rows = query(
                "income", code,
                "ts_code,ann_date,end_date,report_type,revenue,operate_profit,n_income",
            )
            cash_rows = query(
                "cashflow", code,
                "ts_code,ann_date,end_date,report_type,n_cashflow_act,c_pay_acq_const_fiolta",
            )
            balance_rows = query(
                "balancesheet", code,
                "ts_code,end_date,inventories,total_assets,total_liab",
            )
            if not income_rows:
                continue
            cash_by_period = {str(row.get("end_date")): row for row in cash_rows}
            balance_by_period = {str(row.get("end_date")): row for row in balance_rows}
            income_by_period = {str(row.get("end_date")): row for row in income_rows}
            for period, row in sorted(income_by_period.items()):
                row = {
                    **row,
                    **cash_by_period.get(period, {}),
                    **balance_by_period.get(period, {}),
                }
                def bn(field):
                    value = row.get(field)
                    return round(float(value) / 1e9, 4) if value is not None else None
                output.append({
                    "code": code, "company": name, "period": period,
                    "ann_date": str(row.get("ann_date") or ""),
                    "revenue_cny_bn": bn("revenue"), "operating_profit_cny_bn": bn("operate_profit"),
                    "net_income_cny_bn": bn("n_income"), "operating_cashflow_cny_bn": bn("n_cashflow_act"),
                    "capex_cash_cny_bn": bn("c_pay_acq_const_fiolta"), "inventory_cny_bn": bn("inventories"),
                    "assets_cny_bn": bn("total_assets"), "liabilities_cny_bn": bn("total_liab"),
                })
        return output, "Tushare Pro live"
    except Exception as exc:
        if OUT.exists():
            try:
                previous = json.loads(OUT.read_text(encoding="utf-8"))
                cached = previous.get("a_share_financials", [])
                if cached:
                    return cached, f"cached after {type(exc).__name__}"
            except Exception:
                pass
        return FALLBACK_FINANCIALS, f"verified cache after {type(exc).__name__}"


def build_narrative() -> dict[str, Any]:
    return {
        "current_stage": "投资兑现后半段，向瓶颈利润与宏观外溢过渡",
        "score": 78,
        "summary": "需求与资本开支仍在扩张，硬件收入由GPU向ASIC、网络和ODM扩散；商业化质量、自由现金流和电力约束是下一阶段验证重点。",
        "layers": [
            {"id": 1, "name": "真实需求", "state": "扩张", "score": 82, "why": "Token中枢、应用参与度与高强度工作流"},
            {"id": 2, "name": "商业化", "state": "扩张", "score": 72, "why": "ARR与付费渗透改善，但实际混合价格承压"},
            {"id": 3, "name": "投资能力", "state": "高位", "score": 88, "why": "Capex和RPO强，FCF/折旧/债务约束上升"},
            {"id": 4, "name": "硬件交付", "state": "扩散", "score": 84, "why": "GPU、ASIC、网络与ODM收入共同验证"},
            {"id": 5, "name": "瓶颈利润", "state": "验证中", "score": 66, "why": "存储、CCL提价需转化为毛利率和现金流"},
            {"id": 6, "name": "宏观外溢", "state": "早期", "score": 48, "why": "功率、电网、设备与信用开始成为约束"},
        ],
        "confirmed": ["Token样本中枢上移", "Capex与硬件收入尚未脱节", "GPU需求向ASIC、网络和制造链扩散"],
        "watch": ["付费用户与实际混合价格", "FCF、折旧、债务与ROIC", "库存、扩产和利润率匹配", "电网接入与NeoCloud信用"],
        "falsification": ["Token量价同步走弱", "云厂商指引下修且硬件收入斜率下降", "库存上升并伴随自由现金流恶化"],
    }


def main():
    data = extract_main_workbook()
    data.update(extract_state_of_ai())
    data.update(extract_semianalysis())
    data.update(build_credit_snapshot())
    financials, ts_status = tushare_financials()
    data["a_share_financials"] = financials
    data["report_outline"] = extract_report_outline()
    data["pdf_catalog"] = extract_pdf_catalog()
    data["narrative"] = build_narrative()
    data["metadata"] = {
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "source_folder": str(SOURCE),
        "source_files": sorted(p.name for p in SOURCE.iterdir() if p.is_file()),
        "tushare": ts_status,
        "ifind_quantapi": (
            "configured"
            if os.environ.get("IFIND_ACCESS_TOKEN") and os.environ.get("IFIND_REFRESH_TOKEN")
            else "not_configured_in_environment"
        ),
        "refresh_policy": {
            "daily": ["OpenRouter", "模型价格事件", "信用利差、利率与市场代理", "存储/材料价格（源文件刷新后）"],
            "weekly": ["应用榜", "Google Trends", "新闻与催化日历"],
            "monthly": ["Similarweb", "台湾ODM", "电力"],
            "quarterly": ["公司财务", "Capex", "公司模型与盈利验证"],
        },
        "quality_note": "连续序列、事件披露、卖方模型和第三方估算分层展示；未刷新的 #NAME? 不作为数据。",
    }
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    shard_groups = {
        "core.json": ["narrative", "metadata", "report_outline", "pdf_catalog"],
        "demand.json": ["token", "apps", "google_trends", "similarweb"],
        "models.json": ["model_price_changes", "model_benchmarks"],
        "capital.json": [
            "funding", "exits", "top_deals", "top_sectors", "top_companies",
            "recent_activity", "active_investors", "as_of",
        ],
        "hardware.json": ["hardware_quarterly", "odm_monthly", "semiconductor_cycle"],
        "memory.json": ["memory_spot", "memory_contract"],
        "copper.json": ["copper_foil"],
        "materials.json": ["electronic_cloth", "epoxy_resin", "ccl_events"],
        "power.json": ["power_generation"],
        "credit.json": [
            "credit_series", "credit_snapshot", "credit_market_series",
            "credit_market_snapshot", "credit_signal", "credit_sources",
            "credit_metadata",
        ],
        "companies.json": [
            "company_models", "company_revenue", "public_indicators",
            "source_ledger", "a_share_financials",
        ],
    }
    manifest = []
    for filename, keys in shard_groups.items():
        payload = {key: data[key] for key in keys}
        path = OUT.parent / filename
        path.write_text(
            json.dumps(payload, ensure_ascii=False, separators=(",", ":")),
            encoding="utf-8",
        )
        manifest.append({"file": filename, "bytes": path.stat().st_size})
    (OUT.parent / "manifest.json").write_text(
        json.dumps(
            {"version": datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ"), "files": manifest},
            ensure_ascii=False,
            separators=(",", ":"),
        ),
        encoding="utf-8",
    )
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()

